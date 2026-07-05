"""
AX Sejong Command — Phase 1+2
> 백엔드: Python 3.10+ (FastAPI)
> 프론트엔드: HTML + CSS + Vanilla JS
> 데이터 저장소: Notion API (읽기 전용)
"""

import os
import re
import time
import requests
from datetime import date
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv
from fastapi import FastAPI, Request
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles

load_dotenv()

# ========================================
# 환경 설정
# ========================================

NOTION_API_KEY        = os.getenv("NOTION_API_KEY")
NOTION_PROJECTS_DB_ID = os.getenv("NOTION_PROJECTS_DB_ID")
NOTION_WEEKLY_PAGE_ID   = "1f4963cc-8170-8078-b4de-eca9dec4b465"
MONTHLY_COMPANY_PAGE_ID = "353963cc-8170-8177-835b-deeb188c2a0f"
MONTHLY_SEJONG_PAGE_ID  = "352963cc-8170-8146-bf57-ee1c87d96ba7"
HISTORY_PAGE_ID         = "363963cc-8170-81de-a362-d9294436af0b"
# ── Phase 5 v2.2: 사업히스토리 DB ID (보관용, v3에서 교체)
HISTORY_DB_ID           = "19425ad2-e777-4aa5-88d8-54a8ffb65d76"
# ── Phase 5 v3.0: 신규 DB 2개
HISTORY_V3_DB_ID        = "9d8eb5cd-43a2-4031-949d-bb9b6bc23986"   # 프로젝트별 상세
HISTORY_YEARLY_DB_ID    = "5c976da8-e03c-42e1-8134-b359d762efec"   # 연도별합계 (원 단위)

# ── Phase 6: 조직도 DB ID
STAFF_DB_ID   = "7d93242628714413b4c31bffbadd3e03"
PROJECT_DB_ID = "e22b899799494b4e9a9edb8521d2084f"

# ── Phase 7: 영업보고 페이지 ID
SALES_PAGE_ID = "369963cc-8170-816e-b723-efc00c45e3fe"

# ── Phase 8: RISK 알림 페이지 ID
RISK_PAGE_ID = "369963cc-8170-81bd-adc2-d4340f24a2d6"

# ── Phase 9: 세미나 알림 페이지 ID
SEMINAR_PAGE_ID = "369963cc-8170-815e-9ca9-ff3ea7f9d2fb"

NOTION_HEADERS = {
    "Authorization": f"Bearer {NOTION_API_KEY}",
    "Notion-Version": "2022-06-28",
}

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# ========================================
# 키워드 상수
# ========================================

ISSUE_KEYWORDS   = ["오류", "실패", "미작동", "장애", "중단", "이슈", "오작동", "미재현"]
RISK_KEYWORDS    = ["지연", "집중 관리 필요", "확인 필요", "모니터링 필요", "우려", "주의", "시급", "차질"]
DEADLINE_PATTERN = r"[~～]\s*(\d{1,2})/(\d{1,2})"

# ========================================
# 캐시 (메모리) — 10분간 유지
# ========================================

_weekly_cache = {
    "top_blocks": None,   # 캐시된 최상위 블록
    "timestamp":  0,      # 저장 시각 (time.time())
}
CACHE_TTL = 600  # 10분

_monthly_company_cache = {"top_blocks": None, "timestamp": 0}
_monthly_sejong_cache  = {"top_blocks": None, "timestamp": 0}
_history_cache         = {"records": None, "timestamp": 0}
_history_v3_cache      = {"records": None, "timestamp": 0}   # v3.0 프로젝트별
_history_yearly_cache  = {"records": None, "timestamp": 0}   # v3.0 연도별합계
_org_cache             = {"data": None, "timestamp": 0}   # Phase 6: 조직도 캐시
_sales_cache           = {"data": None, "timestamp": 0}   # Phase 7: 영업보고 캐시
_risk_cache            = {"data": None, "timestamp": 0}   # Phase 8: RISK 알림 캐시
_seminar_cache         = {"data": None, "timestamp": 0}   # Phase 9: 세미나 알림 캐시


# ========================================
# PHASE 1 — 세종사업현황 (수정 금지)
# ========================================

def get_notion_projects() -> list:
    """노션 DB에서 세종사업현황 목록을 가져옵니다."""
    url     = f"https://api.notion.com/v1/databases/{NOTION_PROJECTS_DB_ID}/query"
    payload = {"sorts": [{"property": "매출액", "direction": "descending"}]}

    try:
        res = requests.post(url, headers=NOTION_HEADERS, json=payload, timeout=10)
        res.raise_for_status()
        results = res.json().get("results", [])

        projects = []
        for item in results:
            props = item.get("properties", {})

            def get_title(key):
                try: return props[key]["title"][0]["plain_text"]
                except: return ""

            def get_text(key):
                try: return props[key]["rich_text"][0]["plain_text"]
                except: return ""

            def get_number(key):
                try: return props[key]["number"]
                except: return None

            def get_select(key):
                try: return props[key]["select"]["name"]
                except: return ""

            name = get_title("사업명")
            if not name:
                continue

            projects.append({
                "name":    name,
                "client":  get_text("고객사"),
                "pm":      get_text("PM"),
                "revenue": get_number("매출액"),
                "period":  get_text("계약기간"),
                "status":  get_select("상태"),
                "target":  get_text("운영대상"),
                "summary": get_text("사업요약"),
            })
        return projects

    except Exception as e:
        print(f"[ERROR] 노션 프로젝트 API 에러: {e}")
        return []


# ========================================
# PHASE 2 — PM 주간회의
# ========================================

def fetch_block_children(block_id: str) -> list:
    """특정 블록의 자식 블록 전체를 가져옵니다 (페이지네이션 처리)."""
    url      = f"https://api.notion.com/v1/blocks/{block_id}/children"
    children = []
    cursor   = None

    while True:
        params = {"page_size": 100}
        if cursor:
            params["start_cursor"] = cursor
        try:
            res = requests.get(url, headers=NOTION_HEADERS, params=params, timeout=10)
            res.raise_for_status()
            data = res.json()
            children.extend(data.get("results", []))
            if not data.get("has_more"):
                break
            cursor = data.get("next_cursor")
        except Exception as e:
            print(f"[ERROR] 자식 블록 로드 실패 ({block_id}): {e}")
            break

    return children


def fetch_block_children_all(block_id: str) -> list:
    """
    특정 블록의 모든 자식 블록을 재귀적으로 가져옵니다.
    1단계 자식뿐만 아니라, 하위 자식들도 모두 가져옵니다.
    """
    # 1단계 자식 가져오기
    children = fetch_block_children(block_id)
    all_children = list(children)  # 복사
    
    # 각 자식 블록에 대해 재귀적으로 하위 자식 가져오기
    for child in children:
        if child.get("has_children", False):
            # 하위 자식들을 재귀적으로 가져와서 합치기
            grandchildren = fetch_block_children_all(child["id"])
            all_children.extend(grandchildren)
    
    return all_children



def get_block_text(block: dict) -> str:
    """블록에서 plain_text를 추출합니다."""
    btype = block.get("type", "")
    rich  = block.get(btype, {}).get("rich_text", [])
    return "".join(t.get("plain_text", "") for t in rich)


def is_bold_block(block: dict) -> bool:
    """블록 텍스트가 볼드인지 확인합니다."""
    btype = block.get("type", "")
    rich  = block.get(btype, {}).get("rich_text", [])
    return any(t.get("annotations", {}).get("bold") for t in rich)


def get_top_blocks(force_refresh: bool = False) -> list:
    """
    노션 주간회의 페이지의 최상위 블록을 가져옵니다.
    캐시가 유효하면 캐시를 반환하고, 만료되면 노션 API를 호출합니다.
    """
    global _weekly_cache
    now = time.time()

    if (not force_refresh
            and _weekly_cache["top_blocks"] is not None
            and now - _weekly_cache["timestamp"] < CACHE_TTL):
        print("[CACHE] 캐시에서 블록 반환")
        return _weekly_cache["top_blocks"]

    print("[API] 노션에서 최상위 블록 새로 가져오기")
    blocks = fetch_block_children(NOTION_WEEKLY_PAGE_ID)
    _weekly_cache["top_blocks"] = blocks
    _weekly_cache["timestamp"]  = now
    return blocks


def get_weeks(top_blocks: list) -> list:
    """최상위 블록에서 주차 목록을 추출합니다."""
    weeks = []
    for block in top_blocks:
        if block.get("type") == "heading_2":
            text = get_block_text(block)
            if "[26-" in text or "주차" in text:
                weeks.append(text)
    return weeks


def get_projects_for_week(top_blocks: list, week_index: int) -> tuple:
    """
    특정 주차의 사업 목록과 분석 섹션을 파싱합니다.

    핵심 로직:
    1. 최상위 블록에서 주차 범위를 자른다
    2. heading_3 으로 사업 섹션을 구분한다
    3. 볼드 불릿("진행사항"/"특이사항"/"이슈사항") ID를 수집한다
    4. 해당 ID의 자식 블록을 ThreadPoolExecutor로 병렬 fetch한다
    5. 자식 블록 내용을 각 사업 항목에 매핑한다
    6. ※ 로 시작하는 heading_3 이후를 분석 섹션으로 파싱한다

    평탄화(flatten)를 하지 않으므로 중복 삽입 버그가 없습니다.
    """

    # ── 주차 범위 자르기 ──────────────────────────────────
    positions = [i for i, b in enumerate(top_blocks)
                 if b.get("type") == "heading_2"
                 and ("[26-" in get_block_text(b) or "주차" in get_block_text(b))]

    if not positions or week_index >= len(positions):
        return [], []

    s = positions[week_index]
    e = positions[week_index + 1] if week_index + 1 < len(positions) else len(top_blocks)
    week_blocks = top_blocks[s:e]

    # ── heading_3 위치 파악 ───────────────────────────────
    project_headings = []   # (블록인덱스, 텍스트)
    analysis_index   = None

    for i, block in enumerate(week_blocks):
        if block.get("type") != "heading_3":
            continue
        text = get_block_text(block).strip()
        if text.startswith("※"):
            # ※ 로 시작하면 분석 섹션 (AI과제 사업명과 구분)
            analysis_index = i
        elif re.match(r"^\d+\.", text):
            project_headings.append((i, text))

    # ── 각 사업의 볼드 불릿 ID 수집 ─────────────────────
    # section_bullets: {block_id: {"section": "progress"|"special"|"issues", "proj_idx": int}}
    section_bullets = {}

    for proj_idx, (heading_i, _) in enumerate(project_headings):
        # 다음 heading_3 직전까지
        next_i = (project_headings[proj_idx + 1][0]
                  if proj_idx + 1 < len(project_headings)
                  else (analysis_index or len(week_blocks)))

        for block in week_blocks[heading_i + 1:next_i]:
            if (block.get("type") == "bulleted_list_item"
                    and is_bold_block(block)
                    and block.get("has_children")):
                text = get_block_text(block).strip()
                bid  = block["id"]
                if "진행사항" in text:
                    section_bullets[bid] = {"section": "progress", "proj_idx": proj_idx}
                elif "특이사항" in text:
                    section_bullets[bid] = {"section": "special",  "proj_idx": proj_idx}
                elif "이슈사항" in text:
                    section_bullets[bid] = {"section": "issues",   "proj_idx": proj_idx}

    # ── 볼드 불릿 자식 블록 병렬 fetch ───────────────────
    children_map = {}   # {block_id: [child_blocks]}
    if section_bullets:
        with ThreadPoolExecutor(max_workers=8) as pool:
            futures = {pool.submit(fetch_block_children_all, bid): bid
                       for bid in section_bullets}
            for future in as_completed(futures):
                bid = futures[future]
                try:
                    children_map[bid] = future.result()
                except Exception as e:
                    print(f"[ERROR] 병렬 fetch 실패 ({bid}): {e}")
                    children_map[bid] = []

    # ── 사업 데이터 조립 ──────────────────────────────────
    projects = []
    for proj_idx, (heading_i, heading_text) in enumerate(project_headings):
        m = re.match(r"^\d+\.\s+(.+?)\s*[\(（](.+?)[\)）]", heading_text.strip())
        if not m:
            continue

        # 노션 원본 사업명/PM 그대로 사용 (요약·축약 없음)
        project = {
            "name":           m.group(1).strip(),
            "pm":             m.group(2).strip(),
            "progress":       [],
            "special":        [],
            "issues":         [],
            "status":         "정상",
            "issue_count":    0,
            "risk_count":     0,
            "deadline_count": 0,
            "has_special":    False,
        }

        # 자식 블록에서 항목 수집 (이 사업에 해당하는 것만)
        for bid, meta in section_bullets.items():
            if meta["proj_idx"] != proj_idx:
                continue
            for kid in children_map.get(bid, []):
                if kid.get("type") != "bulleted_list_item":
                    continue
                text = get_block_text(kid).strip()
                if not text or is_bold_block(kid):
                    continue
                section = meta["section"]
                project[section].append(text)
                if section == "special":
                    project["has_special"] = True

        projects.append(project)

    # ── 키워드 감지 ───────────────────────────────────────
    projects = [detect_keywords(p) for p in projects]

    # ── 분석 섹션 파싱 ────────────────────────────────────
    analysis = []
    if analysis_index is not None:
        for block in week_blocks[analysis_index + 1:]:
            if block.get("type") == "heading_3":
                break
            if block.get("type") == "bulleted_list_item":
                text = get_block_text(block).strip()
                if text and not is_bold_block(block):
                    analysis.append(text)

    return projects, analysis


def detect_keywords(project: dict) -> dict:
    """키워드 감지 및 상태 판정."""
    today    = date.today()
    all_text = " ".join(project["progress"] + project["special"] + project["issues"])

    issue_cnt    = sum(1 for kw in ISSUE_KEYWORDS if kw in all_text)
    risk_cnt     = sum(1 for kw in RISK_KEYWORDS  if kw in all_text)
    deadline_cnt = 0

    for m in re.finditer(DEADLINE_PATTERN, all_text):
        try:
            d    = date(today.year, int(m.group(1)), int(m.group(2)))
            diff = (d - today).days
            if 0 <= diff <= 7:
                deadline_cnt += 1
        except Exception:
            pass

    project["issue_count"]    = issue_cnt
    project["risk_count"]     = risk_cnt
    project["deadline_count"] = deadline_cnt
    project["status"] = (
        "이슈"    if issue_cnt    > 0 else
        "위험"    if risk_cnt     > 0 else
        "마감임박" if deadline_cnt > 0 else
        "정상"
    )
    return project


def calculate_kpi(projects: list) -> dict:
    return {
        "total":    len(projects),
        "issue":    sum(1 for p in projects if p["issue_count"]    > 0),
        "risk":     sum(1 for p in projects if p["risk_count"]     > 0),
        "special":  sum(1 for p in projects if p["has_special"]),
        "deadline": sum(1 for p in projects if p["deadline_count"] > 0),
    }


def get_weekly_data(week_index: int = 0, force_refresh: bool = False) -> dict:
    """주간회의 데이터를 반환합니다."""
    top_blocks = get_top_blocks(force_refresh=force_refresh)
    weeks      = get_weeks(top_blocks)

    if not weeks:
        return {"success": False, "error": "주차 데이터 없음", "data": {}}
    if week_index >= len(weeks):
        return {"success": False, "error": f"주차 인덱스 초과 (총 {len(weeks)}주차)", "data": {}}

    projects, analysis = get_projects_for_week(top_blocks, week_index)
    kpi = calculate_kpi(projects)

    return {
        "success": True,
        "data": {
            "week_title": weeks[week_index],
            "weeks":      weeks,
            "projects":   projects,
            "analysis":   analysis,
            "kpi":        kpi,
        }
    }


# ========================================
# PHASE 3 — 회사 월간보고
# ========================================

def fetch_table_rows(table_block: dict) -> list:
    """table 블록의 자식(table_row) 블록을 가져와 2차원 배열로 반환합니다."""
    if not table_block.get("has_children"):
        return []
    children = fetch_block_children(table_block["id"])
    rows = []
    for row_block in children:
        if row_block.get("type") == "table_row":
            cells = []
            for cell in row_block.get("table_row", {}).get("cells", []):
                text = "".join(t.get("plain_text", "") for t in cell)
                cells.append(text.strip())
            if any(cells):
                rows.append(cells)
    return rows


def extract_number_from_text(text: str, keyword: str) -> int:
    """텍스트에서 키워드 뒤의 숫자를 추출합니다. 예: '진행 프로젝트 : 6건' → 6"""
    if keyword not in text:
        return 0
    after = text[text.index(keyword) + len(keyword):]
    m = re.search(r"\d+", after)
    return int(m.group()) if m else 0


def get_months_company(top_blocks: list) -> list:
    """heading_2에서 월 목록을 추출합니다."""
    months = []
    for block in top_blocks:
        if block.get("type") == "heading_2":
            text = get_block_text(block)
            if "[26-" in text and "월" in text:
                months.append(text)
    return months


def get_month_blocks_company(top_blocks: list, month_index: int = 0) -> list:
    """특정 월의 블록만 잘라냅니다."""
    positions = [i for i, b in enumerate(top_blocks)
                 if b.get("type") == "heading_2"
                 and "[26-" in get_block_text(b)]

    if not positions or month_index >= len(positions):
        return []

    s = positions[month_index]
    e = positions[month_index + 1] if month_index + 1 < len(positions) else len(top_blocks)
    return top_blocks[s:e]


def parse_month_company(month_blocks: list) -> dict:
    """회사 월간보고 한 달치 블록을 파싱합니다."""
    result = {
        "month_title": "",
        "kpi":         {"total": 0, "issue": 0, "meeting": 0},
        "workforce":   [],
        "issue_table": [],
        "meetings":    [],
        "analysis":    [],
    }

    in_meeting_section = False
    in_analysis        = False
    table_count        = 0

    for block in month_blocks:
        btype = block.get("type", "")
        text  = get_block_text(block).strip()

        # ── 월 제목
        if btype == "heading_2" and "[26-" in text:
            result["month_title"] = text

        # ── 분석 섹션 감지
        elif btype == "heading_3" and text.startswith("※"):
            in_analysis        = True
            in_meeting_section = False

        # ── 분석 내용 수집
        elif in_analysis and btype == "bulleted_list_item" and text:
            result["analysis"].append(text)

        # ── KPI 수치 추출
        elif btype == "paragraph" and "진행 프로젝트 :" in text:
            result["kpi"]["total"] = extract_number_from_text(text, "진행 프로젝트 :")

        elif btype == "paragraph" and "이슈 프로젝트현황 :" in text:
            result["kpi"]["issue"] = extract_number_from_text(text, "이슈 프로젝트현황 :")

        # ── 미팅 섹션 진입
        elif btype == "paragraph" and ("03." in text or "사업 확산" in text):
            in_meeting_section = True

        # ── 미팅 내용 수집
        elif in_meeting_section and btype == "paragraph" and text and not is_bold_block(block):
            if re.match(r"^0[4-9]\.", text) or "인력투입" in text:
                in_meeting_section = False
            else:
                result["meetings"].append(text)
                result["kpi"]["meeting"] += 1

        # ── 테이블 파싱
        elif btype == "table":
            rows = fetch_table_rows(block)
            table_count += 1

            if table_count == 1:
                # 첫 번째 테이블: 인력투입현황
                if len(rows) >= 2:
                    headers = rows[0]
                    for row in rows[1:]:
                        if row and "세종개발본부" in row[0]:
                            for i, h in enumerate(headers[1:], 1):
                                if i < len(row):
                                    result["workforce"].append({"label": h, "value": row[i]})

            elif table_count == 2:
                # 두 번째 테이블: 이슈사항
                if len(rows) >= 2:
                    for row in rows[1:]:
                        if len(row) >= 6 and any(row):
                            result["issue_table"].append({
                                "client":     row[0] if len(row) > 0 else "",
                                "project":    row[1] if len(row) > 1 else "",
                                "budget":     row[2] if len(row) > 2 else "",
                                "period":     row[3] if len(row) > 3 else "",
                                "pm":         row[4] if len(row) > 4 else "",
                                "issue":      row[5] if len(row) > 5 else "",
                                "resolution": row[6] if len(row) > 6 else "",
                            })

    return result


def get_monthly_data_company(month_index: int = 0, force_refresh: bool = False) -> dict:
    """회사 월간보고 데이터를 반환합니다."""
    global _monthly_company_cache
    now = time.time()

    if (not force_refresh
            and _monthly_company_cache["top_blocks"] is not None
            and now - _monthly_company_cache["timestamp"] < CACHE_TTL):
        top_blocks = _monthly_company_cache["top_blocks"]
    else:
        top_blocks = fetch_block_children(MONTHLY_COMPANY_PAGE_ID)
        _monthly_company_cache = {"top_blocks": top_blocks, "timestamp": now}

    months       = get_months_company(top_blocks)
    month_blocks = get_month_blocks_company(top_blocks, month_index)

    if not months:
        return {"success": False, "error": "월 데이터 없음", "data": {}}

    data           = parse_month_company(month_blocks)
    data["months"] = months

    return {"success": True, "data": data}


# ========================================
# PHASE 4 — 세종 월간분석
# ========================================

def get_months_sejong(top_blocks: list) -> list:
    """heading_2에서 월 목록을 추출합니다."""
    months = []
    for block in top_blocks:
        if block.get("type") == "heading_2":
            text = get_block_text(block)
            if "[26-" in text and "월" in text:
                months.append(text)
    return months


def get_month_blocks_sejong(top_blocks: list, month_index: int = 0) -> list:
    """특정 월의 블록만 잘라냅니다."""
    positions = [i for i, b in enumerate(top_blocks)
                 if b.get("type") == "heading_2"
                 and "[26-" in get_block_text(b)]

    if not positions or month_index >= len(positions):
        return []

    s = positions[month_index]
    e = positions[month_index + 1] if month_index + 1 < len(positions) else len(top_blocks)
    return top_blocks[s:e]


def parse_month_sejong(month_blocks: list) -> dict:
    """세종 월간분석 한 달치 블록을 파싱합니다."""
    result = {
        "month_title": "",
        "period":      "",
        "kpi_table":   [],
        "projects":    [],
        "analysis":    [],
    }

    cur_project  = None
    cur_section  = None
    in_kpi       = False
    in_projects  = False
    in_analysis  = False

    for block in month_blocks:
        btype = block.get("type", "")
        text  = get_block_text(block).strip()

        # ── 월 제목
        if btype == "heading_2" and "[26-" in text:
            result["month_title"] = text

        # ── 기간/작성자 paragraph
        elif btype == "paragraph" and "기간:" in text:
            result["period"] = text

        # ── heading_3 감지
        elif btype == "heading_3":
            prev_in_projects = in_projects
            in_kpi      = "KPI" in text or "1." in text
            in_projects = "사업별" in text or "2." in text
            in_analysis = text.startswith("※")

            # 분석 섹션 진입 시 마지막 사업 저장
            if in_analysis and cur_project:
                result["projects"].append(cur_project)
                cur_project = None

        # ── KPI 테이블 (heading_3 "1. KPI 요약" 다음에 오는 table)
        elif in_kpi and btype == "table":
            rows = fetch_table_rows(block)
            for row in rows[1:]:  # 헤더 제외
                if len(row) >= 2 and row[0]:
                    result["kpi_table"].append({"label": row[0], "value": row[1] if len(row) > 1 else ""})
            in_kpi = False

        # ── heading_4: 사업 섹션 시작
        elif btype == "heading_4" and in_projects:
            if cur_project:
                result["projects"].append(cur_project)

            m = re.match(r"^\d+\)\s+(.+?)\s*[\(（](.+?)[\)）]\s*$", text)
            cur_project = {
                "name":        m.group(1).strip() if m else text,
                "pm":          m.group(2).strip() if m else "",
                "completed":   [],
                "in_progress": [],
                "issues":      [],
            }
            cur_section = None

        # ── 볼드 불릿: 섹션 구분자(완료/진행/이슈) 또는 분석 항목
        elif btype == "bulleted_list_item" and is_bold_block(block) and text:
            if in_analysis:
                # ※ 분석 섹션의 볼드 항목은 분석 내용으로 수집
                result["analysis"].append(text)
            elif cur_project:
                # 프로젝트 섹션 구분자
                if "완료" in text:
                    cur_section = "completed"
                elif "진행" in text:
                    cur_section = "in_progress"
                elif "이슈" in text or "특이" in text:
                    cur_section = "issues"

        # ── 일반 불릿: 실제 항목 수집
        elif btype == "bulleted_list_item" and not is_bold_block(block) and text:
            if in_analysis:
                result["analysis"].append(text)
            elif cur_project and cur_section:
                cur_project[cur_section].append(text)

    # 마지막 사업 저장
    if cur_project:
        result["projects"].append(cur_project)

    return result


def expand_bold_bullet_children(blocks: list) -> list:
    """
    볼드 불릿(완료 항목/진행 중/이슈)의 자식 블록을 병렬로 가져와 평탄화합니다.

    실제 Notion 구조:
      heading_4 (has_children=False)
      bulleted_list_item bold "완료 항목" (has_children=True)  ← 자식 fetch 필요
        → 실제 완료 항목들
      bulleted_list_item bold "진행 중" (has_children=True)
        → 실제 진행 항목들
    """
    bold_ids = [
        b["id"] for b in blocks
        if (b.get("type") == "bulleted_list_item"
            and is_bold_block(b)
            and b.get("has_children"))
    ]

    if not bold_ids:
        return blocks

    children_map = {}
    with ThreadPoolExecutor(max_workers=8) as pool:
        futures = {pool.submit(fetch_block_children, bid): bid for bid in bold_ids}
        for future in as_completed(futures):
            bid = futures[future]
            try:
                children_map[bid] = future.result()
            except Exception as e:
                print(f"[ERROR] 볼드 불릿 자식 fetch 실패 ({bid}): {e}")
                children_map[bid] = []

    flat = []
    for block in blocks:
        flat.append(block)
        if (block.get("type") == "bulleted_list_item"
                and is_bold_block(block)
                and block["id"] in children_map):
            flat.extend(children_map[block["id"]])

    return flat


def get_monthly_data_sejong(month_index: int = 0, force_refresh: bool = False) -> dict:
    """세종 월간분석 데이터를 반환합니다."""
    global _monthly_sejong_cache
    now = time.time()

    if (not force_refresh
            and _monthly_sejong_cache["top_blocks"] is not None
            and now - _monthly_sejong_cache["timestamp"] < CACHE_TTL):
        top_blocks = _monthly_sejong_cache["top_blocks"]
    else:
        top_blocks = fetch_block_children(MONTHLY_SEJONG_PAGE_ID)
        _monthly_sejong_cache = {"top_blocks": top_blocks, "timestamp": now}

    months       = get_months_sejong(top_blocks)
    month_blocks = get_month_blocks_sejong(top_blocks, month_index)

    if not months:
        return {"success": False, "error": "월 데이터 없음", "data": {}}

    # 볼드 불릿(완료/진행/이슈) 자식 블록 평탄화
    month_blocks = expand_bold_bullet_children(month_blocks)

    data           = parse_month_sejong(month_blocks)
    data["months"] = months

    return {"success": True, "data": data}


# ========================================
# PHASE 5 — 세종사업히스토리
# ========================================

def parse_number(text: str) -> int:
    """
    숫자 문자열에서 정수를 추출합니다.
    예: "5,030,248,818" → 5030248818
        "-109,641,969"  → -109641969
        "-"             → 0
    """
    if not text or text.strip() in ["-", ""]:
        return 0
    # 콤마·공백 등 제거, 숫자와 마이너스만 남김
    cleaned = re.sub(r"[^\d\-]", "", text.strip())
    try:
        return int(cleaned)
    except Exception:
        return 0


def fetch_history_db(force_refresh: bool = False) -> list:
    """노션 DB에서 사업히스토리 전체 레코드를 조회합니다. (v2.2)"""
    global _history_cache
    now = time.time()

    if (not force_refresh
            and _history_cache["records"] is not None
            and now - _history_cache["timestamp"] < CACHE_TTL):
        print("[CACHE] 히스토리 DB 캐시 반환")
        return _history_cache["records"]

    print("[API] 히스토리 DB 데이터 새로 가져오기")
    url     = f"https://api.notion.com/v1/databases/{HISTORY_DB_ID}/query"
    records = []
    body    = {"page_size": 100}

    while True:
        try:
            res = requests.post(url, headers=NOTION_HEADERS, json=body, timeout=10)
            res.raise_for_status()
            data = res.json()
            for page in data.get("results", []):
                record = parse_db_record(page)
                if record:
                    records.append(record)
            if not data.get("has_more"):
                break
            body["start_cursor"] = data["next_cursor"]
        except Exception as e:
            print(f"[ERROR] 히스토리 DB 조회 실패: {e}")
            break

    _history_cache = {"records": records, "timestamp": now}
    return records


def parse_db_record(page: dict) -> dict | None:
    """노션 DB 레코드를 내부 dict로 변환합니다. (v2.2)"""
    props = page.get("properties", {})

    def get_text(key):
        p = props.get(key, {})
        if p.get("type") == "title":
            return "".join(t["plain_text"] for t in p.get("title", []))
        if p.get("type") == "rich_text":
            return "".join(t["plain_text"] for t in p.get("rich_text", []))
        return ""

    def get_number(key):
        p = props.get(key, {})
        return p.get("number") or 0

    def get_select(key):
        p = props.get(key, {})
        sel = p.get("select")
        return sel.get("name", "") if sel else ""

    project_type = get_select("사업구분")
    name         = get_text("사업명")
    if not name:
        return None

    is_excluded = project_type == "제외"

    # 웹앱 자동계산 필드
    revenue    = get_number("매출액")
    cost       = get_number("매출원가")
    op_profit  = get_number("영업이익")
    op_plan    = get_number("영업이익_계획")

    gross_profit = revenue - cost
    gross_rate   = f"{gross_profit / revenue * 100:.1f}%" if revenue else "—"
    op_rate      = f"{op_profit / revenue * 100:.1f}%" if revenue else "—"
    vs_plan      = op_profit - op_plan

    return {
        "project_type": project_type,
        "name":         name,
        "year":         get_select("년도"),
        "client":       get_text("고객사"),
        "period":       get_text("기간"),
        "revenue":      revenue,
        "cost":         cost,
        "gross_profit": gross_profit,
        "gross_rate":   gross_rate,
        "overhead":     get_number("공통비"),
        "op_profit":    op_profit,
        "op_rate":      op_rate,
        "vs_plan":      vs_plan,
        "is_total":     False,
        "is_excluded":  is_excluded,
    }


def build_chart_data(projects: list) -> dict:
    """
    연도별 사업 목록에서 총매출/운영매출/운영외매출을 집계합니다. (v2.1)
    - 운영매출: project_type == "운영"
    - 운영외매출: project_type != "운영" (제외 항목은 둘 다 미포함)
    """
    total = ops = non_ops = 0
    for p in projects:
        if p.get("is_total") or p.get("is_excluded"):
            continue
        rev = p.get("revenue", 0) or 0
        total += rev
        if p.get("project_type") == "운영":
            ops += rev
        else:
            non_ops += rev
    return {"total": total, "ops": ops, "non_ops": non_ops}


def calc_year_total(projects: list) -> dict:
    """
    제외 항목을 빼고 합계를 재계산합니다. (v2.1)
    노션의 합계 행은 제외 처리된 항목을 반영하지 못하므로 웹앱이 직접 계산합니다.
    """
    revenue = cost = gross_profit = overhead = op_profit = vs_plan = 0
    count = 0
    for p in projects:
        if p.get("is_total") or p.get("is_excluded"):
            continue
        revenue      += p.get("revenue", 0)      or 0
        cost         += p.get("cost", 0)         or 0
        gross_profit += p.get("gross_profit", 0) or 0
        overhead     += p.get("overhead", 0)     or 0
        op_profit    += p.get("op_profit", 0)    or 0
        vs_plan      += p.get("vs_plan", 0)      or 0
        count        += 1

    gross_rate = f"{gross_profit / revenue * 100:.1f}%" if revenue else "—"
    op_rate    = f"{op_profit    / revenue * 100:.1f}%" if revenue else "—"

    return {
        "project_type": "",
        "name":         f"합계 ({count}건)",
        "client":       "",
        "period":       "",
        "revenue":      revenue,
        "cost":         cost,
        "gross_profit": gross_profit,
        "gross_rate":   gross_rate,
        "overhead":     overhead,
        "op_profit":    op_profit,
        "op_rate":      op_rate,
        "vs_plan":      vs_plan,
        "is_total":     True,
        "is_excluded":  False,
    }


# ── v3.0: 프로젝트별 DB (v3) 조회
def fetch_history_v3_db(force_refresh: bool = False) -> list:
    """[AX]세종사업히스토리_v3에서 프로젝트별 레코드를 조회합니다."""
    global _history_v3_cache
    now = time.time()
    if (not force_refresh
            and _history_v3_cache["records"] is not None
            and now - _history_v3_cache["timestamp"] < CACHE_TTL):
        return _history_v3_cache["records"]

    url     = f"https://api.notion.com/v1/databases/{HISTORY_V3_DB_ID}/query"
    records = []
    body    = {"page_size": 100}
    while True:
        try:
            res = requests.post(url, headers=NOTION_HEADERS, json=body, timeout=10)
            res.raise_for_status()
            data = res.json()
            for page in data.get("results", []):
                record = parse_v3_record(page)
                if record:
                    records.append(record)
            if not data.get("has_more"):
                break
            body["start_cursor"] = data["next_cursor"]
        except Exception as e:
            print(f"[ERROR] 히스토리 v3 DB 조회 실패: {e}")
            break
    _history_v3_cache = {"records": records, "timestamp": now}
    return records


def parse_v3_record(page: dict) -> dict | None:
    """v3 프로젝트 DB 레코드를 파싱합니다. 매출이익/매출이익률은 노션값 그대로 사용."""
    props = page.get("properties", {})

    def get_text(key):
        p = props.get(key, {})
        if p.get("type") == "title":
            return "".join(t["plain_text"] for t in p.get("title", []))
        if p.get("type") == "rich_text":
            return "".join(t["plain_text"] for t in p.get("rich_text", []))
        return ""

    def get_number(key):
        p = props.get(key, {})
        return p.get("number") or 0

    def get_select(key):
        p = props.get(key, {})
        sel = p.get("select")
        return sel.get("name", "") if sel else ""

    name = get_text("사업명")
    if not name:
        return None

    project_type = get_select("사업구분")
    is_excluded  = project_type == "제외"
    revenue      = get_number("매출액")
    cost         = get_number("매출원가")
    gross_profit = get_number("매출이익")
    op_profit    = get_number("영업이익")

    raw_gr = props.get("매출이익률", {}).get("number")
    gross_rate = f"{raw_gr * 100:.1f}%" if raw_gr is not None else ("—" if not revenue else f"{gross_profit / revenue * 100:.1f}%")
    raw_or = props.get("영업이익률", {}).get("number")
    op_rate = f"{raw_or * 100:.1f}%" if raw_or is not None else ("—" if not revenue else f"{op_profit / revenue * 100:.1f}%")

    vs_plan = get_number("계획대비실적") or (op_profit - get_number("영업이익_계획"))

    return {
        "project_type": project_type,
        "name":         name,
        "year":         get_select("년도"),
        "client":       get_text("고객사"),
        "period":       get_text("기간"),
        "revenue":      revenue,
        "cost":         cost,
        "gross_profit": gross_profit,
        "gross_rate":   gross_rate,
        "overhead":     get_number("공통비"),
        "op_profit":    op_profit,
        "op_rate":      op_rate,
        "vs_plan":      vs_plan,
        "is_total":     False,
        "is_excluded":  is_excluded,
    }


# ── v3.0: 연도별합계 DB 조회 (원 단위 주의)
def fetch_history_yearly_db(force_refresh: bool = False) -> list:
    """[AX]세종사업히스토리_연도별합계_v2에서 연간 합계를 조회합니다. 단위: 원."""
    global _history_yearly_cache
    now = time.time()
    if (not force_refresh
            and _history_yearly_cache["records"] is not None
            and now - _history_yearly_cache["timestamp"] < CACHE_TTL):
        return _history_yearly_cache["records"]

    url     = f"https://api.notion.com/v1/databases/{HISTORY_YEARLY_DB_ID}/query"
    records = []
    body    = {"page_size": 100}
    while True:
        try:
            res = requests.post(url, headers=NOTION_HEADERS, json=body, timeout=10)
            res.raise_for_status()
            data = res.json()
            for page in data.get("results", []):
                record = parse_yearly_record(page)
                if record:
                    records.append(record)
            if not data.get("has_more"):
                break
            body["start_cursor"] = data["next_cursor"]
        except Exception as e:
            print(f"[ERROR] 연도별합계 DB 조회 실패: {e}")
            break
    _history_yearly_cache = {"records": records, "timestamp": now}
    return records


def parse_yearly_record(page: dict) -> dict | None:
    """연도별합계 DB 레코드를 파싱합니다. 금액은 원 단위 그대로 반환."""
    props = page.get("properties", {})

    def get_title(key):
        p = props.get(key, {})
        if p.get("type") == "title":
            return "".join(t["plain_text"] for t in p.get("title", []))
        return ""

    def get_number(key):
        p = props.get(key, {})
        return p.get("number") or 0

    def get_select(key):
        p = props.get(key, {})
        sel = p.get("select")
        return sel.get("name", "") if sel else ""

    year_sel    = get_select("년도")
    record_type = get_select("유형")
    if not year_sel or not record_type:
        return None

    return {
        "title":          get_title("구분"),
        "year":           year_sel,           # "2026년"
        "type":           record_type,        # "프로젝트합계" | "내부원가현행화"
        "revenue":        get_number("매출액"),       # 원 단위
        "cost":           get_number("매출원가"),
        "gross_profit":   get_number("매출이익"),
        "gross_rate_raw": props.get("매출이익률", {}).get("number"),  # 소수
        "overhead":       get_number("공통비"),
        "op_profit":      get_number("영업이익"),
        "op_rate_raw":    props.get("영업이익률", {}).get("number"),
    }


def get_history_data(force_refresh: bool = False) -> dict:
    """세종사업히스토리 전체 데이터를 반환합니다. (v3.0 — 2-DB 구조)"""

    # 1) v3 프로젝트별 데이터
    v3_records = fetch_history_v3_db(force_refresh)

    # 2) 연도별합계 데이터 (원 단위)
    yearly_records = fetch_history_yearly_db(force_refresh)
    yearly_map = {}  # {"2026": {"내부원가현행화": {...}, "프로젝트합계": {...}}}
    for r in yearly_records:
        yr = r["year"].replace("년", "").strip()
        if yr not in yearly_map:
            yearly_map[yr] = {}
        yearly_map[yr][r["type"]] = r

    # 3) 연도별 프로젝트 그룹핑
    year_order = ["2020년", "2021년", "2022년", "2023년", "2024년", "2025년", "2026년"]
    year_map   = {y: [] for y in year_order}
    for r in v3_records:
        year = r.get("year", "")
        if year in year_map:
            year_map[year].append(r)

    years        = []
    summary_rows = []

    for year_label in year_order:
        year_num = year_label.replace("년", "").strip()
        projects = year_map[year_label]
        yr_data  = yearly_map.get(year_num, {})

        # 연간 합계: 내부원가현행화 우선, 없으면 프로젝트합계 (2020년)
        main_row = yr_data.get("내부원가현행화") or yr_data.get("프로젝트합계")
        ref_row  = yr_data.get("프로젝트합계") if "내부원가현행화" in yr_data else None

        if not main_row and not projects:
            continue

        if main_row:
            rev  = main_row["revenue"]       # 원 단위
            cost = main_row["cost"]
            gp   = main_row["gross_profit"]
            gr_raw = main_row["gross_rate_raw"]
            gr = f"{gr_raw * 100:.1f}%" if gr_raw else "—"
            overhead  = main_row["overhead"]
            op_profit = main_row["op_profit"]
            or_raw    = main_row["op_rate_raw"]
            op_rate   = f"{or_raw * 100:.1f}%" if or_raw else "—"

            summary = {
                "year":         year_num,
                "revenue":      rev,          # 원 단위 (프론트에서 ÷1e8)
                "cost":         cost,
                "gross_profit": gp,
                "gross_rate":   gr,
                "overhead":     overhead,
                "op_profit":    op_profit,
                "op_rate":      op_rate,
            }

            # 프로젝트합계 참고값 (2020년은 ref_row 없음)
            if ref_row:
                ref_gr_raw = ref_row["gross_rate_raw"]
                summary["ref_revenue"]      = ref_row["revenue"]
                summary["ref_cost"]         = ref_row["cost"]
                summary["ref_gross_profit"] = ref_row["gross_profit"]
                summary["ref_gross_rate"]   = f"{ref_gr_raw * 100:.1f}%" if ref_gr_raw else "—"

            summary_rows.append(summary)

        # 프로젝트별 상세 + 차트 데이터 (합계 행 없이)
        years.append({
            "year":     year_num,
            "projects": projects,
            "chart":    build_chart_data(projects),
        })

    # 전년대비 계산 (매출액 + 매출이익 기준, 내부원가현행화끼리 비교)
    for i, s in enumerate(summary_rows):
        if i == 0:
            s["yoy_amount"]    = None
            s["yoy_rate"]      = None
            s["yoy_gp_amount"] = None
            s["yoy_gp_rate"]   = None
        else:
            prev_rev = summary_rows[i - 1]["revenue"]
            curr_rev = s["revenue"]
            diff     = curr_rev - prev_rev
            rate     = diff / prev_rev * 100 if prev_rev else 0
            s["yoy_amount"] = diff
            s["yoy_rate"]   = rate

            prev_gp = summary_rows[i - 1].get("gross_profit") or 0
            curr_gp = s.get("gross_profit") or 0
            gp_diff = curr_gp - prev_gp
            gp_rate = gp_diff / prev_gp * 100 if prev_gp else 0
            s["yoy_gp_amount"] = gp_diff
            s["yoy_gp_rate"]   = gp_rate

    # 손익현황 DB에서 수행조직_매출 가져오기 (KPI용, 단위: 억원)
    try:
        pl_records = fetch_profit_loss_db(force_refresh)
        org_rev_map = {}
        for r in pl_records:
            yr = r.get("연도", "").replace("년", "").strip()
            if yr and r.get("수행조직_매출") is not None:
                org_rev_map[yr] = r["수행조직_매출"]
        for s in summary_rows:
            s["org_revenue"] = org_rev_map.get(s["year"])  # 억원 단위
    except Exception:
        for s in summary_rows:
            s["org_revenue"] = None

    return {
        "success": True,
        "data": {
            "summary": summary_rows,
            "years":   years,
        }
    }


# ========================================
# API 엔드포인트
# ========================================

@app.get("/")
def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/api/projects")
def api_projects():
    """Phase 1: 세종사업현황."""
    try:
        return {"success": True, "data": get_notion_projects()}
    except Exception as e:
        return {"success": False, "error": str(e), "data": []}


@app.get("/api/weekly")
def api_weekly(week: int = 0, refresh: int = 0):
    """
    Phase 2: PM 주간회의.
    ?refresh=1 로 호출하면 캐시를 무시하고 노션에서 새로 가져옵니다.
    새로고침 버튼 클릭 시 refresh=1 을 붙여서 호출하세요.
    """
    try:
        return get_weekly_data(week_index=week, force_refresh=bool(refresh))
    except Exception as e:
        return {"success": False, "error": str(e), "data": {}}


@app.get("/api/weekly/raw")
def api_weekly_raw():
    """디버깅용: 최상위 블록 50개 반환."""
    try:
        blocks = fetch_block_children(NOTION_WEEKLY_PAGE_ID)
        return {"success": True, "data": {"total": len(blocks), "blocks": blocks[:50]}}
    except Exception as e:
        return {"success": False, "error": str(e), "data": {}}

@app.get("/api/weekly/block/{block_id}")
def api_weekly_block(block_id: str):
    """디버깅용: 특정 블록의 자식 블록 반환."""
    try:
        children = fetch_block_children(block_id)
        return {"success": True, "data": {"total": len(children), "children": children}}
    except Exception as e:
        return {"success": False, "error": str(e), "data": {}}


@app.get("/api/monthly")
def api_monthly(month: int = 0, refresh: int = 0):
    """Phase 3: 회사 월간보고."""
    try:
        return get_monthly_data_company(month_index=month, force_refresh=bool(refresh))
    except Exception as e:
        return {"success": False, "error": str(e), "data": {}}


@app.get("/api/monthly-sejong")
def api_monthly_sejong(month: int = 0, refresh: int = 0):
    """Phase 4: 세종 월간분석."""
    try:
        return get_monthly_data_sejong(month_index=month, force_refresh=bool(refresh))
    except Exception as e:
        return {"success": False, "error": str(e), "data": {}}


@app.get("/api/monthly-sejong/debug")
def api_monthly_sejong_debug():
    """디버깅용: 세종 월간분석 블록 구조 확인."""
    try:
        top_blocks   = fetch_block_children(MONTHLY_SEJONG_PAGE_ID)
        months       = get_months_sejong(top_blocks)
        month_blocks = get_month_blocks_sejong(top_blocks, 0)

        h4_blocks = [b for b in month_blocks if b.get("type") == "heading_4"]

        return {
            "success": True,
            "data": {
                "total_top_blocks":  len(top_blocks),
                "months":            months,
                "month_0_blocks":    len(month_blocks),
                "heading4_count":    len(h4_blocks),
                "heading4_list": [
                    {
                        "id":           b["id"],
                        "text":         get_block_text(b),
                        "has_children": b.get("has_children", False),
                    }
                    for b in h4_blocks
                ],
                "first_60_blocks": [
                    {
                        "type":         b.get("type"),
                        "text":         get_block_text(b)[:80],
                        "has_children": b.get("has_children", False),
                        "bold":         is_bold_block(b),
                    }
                    for b in month_blocks[:60]
                ],
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/history")
def api_history(refresh: int = 0):
    """Phase 5: 세종사업히스토리."""
    try:
        return get_history_data(force_refresh=bool(refresh))
    except Exception as e:
        return {"success": False, "error": str(e), "data": {}}


# ========================================
# PHASE 6 — 조직도
# ========================================

def query_notion_db(db_id: str) -> list:
    """노션 DB를 쿼리해서 페이지 목록 반환 (표시순서 오름차순, 페이지네이션 포함)."""
    url  = f"https://api.notion.com/v1/databases/{db_id}/query"
    body = {"sorts": [{"property": "표시순서", "direction": "ascending"}]}
    results = []

    while True:
        try:
            res  = requests.post(url, headers=NOTION_HEADERS, json=body, timeout=10)
            res.raise_for_status()
            data = res.json()
            results.extend(data.get("results", []))
            if not data.get("has_more"):
                break
            body["start_cursor"] = data["next_cursor"]
        except Exception as e:
            print(f"[ERROR] DB 쿼리 실패 ({db_id}): {e}")
            break

    return results


def parse_select(prop: dict) -> str:
    """select 속성 값 추출."""
    try:
        sel = prop.get("select")
        return sel["name"] if sel else ""
    except Exception:
        return ""


def parse_text(prop: dict) -> str:
    """title / rich_text 속성 plain_text 추출."""
    try:
        rt = prop.get("rich_text") or prop.get("title") or []
        return "".join(t["plain_text"] for t in rt)
    except Exception:
        return ""


def parse_email(prop: dict) -> str:
    """email 속성 값 추출."""
    try:
        return prop.get("email") or ""
    except Exception:
        return ""


def parse_phone(prop: dict) -> str:
    """phone_number 속성 값 추출."""
    try:
        return prop.get("phone_number") or ""
    except Exception:
        return ""


def parse_date(prop: dict) -> str:
    """date 속성 start 값 추출."""
    try:
        d = prop.get("date")
        return d["start"] if d else ""
    except Exception:
        return ""


def parse_num(prop: dict) -> int:
    """number 속성 정수 값 추출 (기존 parse_number와 이름 충돌 방지)."""
    try:
        return int(prop.get("number") or 0)
    except Exception:
        return 0


def get_org_data(force_refresh: bool = False) -> dict:
    """사업구조 + 인력 데이터를 노션에서 조합하여 반환."""
    global _org_cache
    now = time.time()

    if (not force_refresh
            and _org_cache["data"] is not None
            and now - _org_cache["timestamp"] < CACHE_TTL):
        print("[CACHE] 조직도 캐시 반환")
        return _org_cache["data"]

    print("[API] 조직도 데이터 새로 가져오기")
    project_pages = query_notion_db(PROJECT_DB_ID)
    staff_pages   = query_notion_db(STAFF_DB_ID)

    # 사업구조 파싱
    projects = []
    for page in project_pages:
        props = page.get("properties", {})
        projects.append({
            "name":        parse_text(props.get("사업명",  {})),
            "color":       parse_text(props.get("색상코드", {})),
            "description": parse_text(props.get("설명",   {})),
            "client":      parse_text(props.get("고객사",  {})),
            "order":       parse_num(props.get("표시순서", {})),
        })

    # 인력 파싱
    staff = []
    for page in staff_pages:
        props = page.get("properties", {})
        staff.append({
            "name":      parse_text(props.get("이름",    {})),
            "grade":     parse_select(props.get("직급",   {})),
            "role":      parse_select(props.get("직책",   {})),
            "project":   parse_select(props.get("소속사업", {})),
            "type":      parse_select(props.get("구분",   {})),
            "join_date": parse_date(props.get("입사일",   {})),
            "birth":     parse_text(props.get("생년월일",  {})),
            "phone":     parse_phone(props.get("전화번호", {})),
            "email1":    parse_email(props.get("이메일1",  {})),
            "email2":    parse_email(props.get("이메일2",  {})),
            "email3":    parse_email(props.get("이메일3",  {})),
            "address":   parse_text(props.get("주소",    {})),
            "photo":     parse_text(props.get("사진파일명", {})),
            "order":     parse_num(props.get("표시순서",  {})),
        })

    result = {"projects": projects, "staff": staff}
    _org_cache["data"]      = result
    _org_cache["timestamp"] = now
    return result


@app.get("/api/org")
def api_org(refresh: int = 0):
    """Phase 6: 조직도 — 사업구조 + 인력 데이터 반환."""
    try:
        return get_org_data(force_refresh=bool(refresh))
    except Exception as e:
        return {"success": False, "error": str(e), "projects": [], "staff": []}


# ========================================
# PHASE 7 — 영업보고
# ========================================

def parse_sales_blocks(top_blocks: list) -> list:
    """
    heading_2 기준으로 주차를 구분하고 각 주차를 파싱합니다.
    - fetch_block_children() / get_block_text() 기존 함수 재사용
    - 노션 페이지 상단이 최신이므로 순서 그대로 반환
    """
    # ── 주차 헤더(heading_2) 위치 수집: "[20XX년" 패턴으로 필터링
    week_positions = [
        i for i, b in enumerate(top_blocks)
        if b.get("type") == "heading_2"
        and re.search(r"\[\d{4}년", get_block_text(b))
    ]

    weeks = []
    for idx, start in enumerate(week_positions):
        end          = week_positions[idx + 1] if idx + 1 < len(week_positions) else len(top_blocks)
        week_blocks  = top_blocks[start:end]

        # ── 주차 라벨 추출 (예: "2026년5월3주차")
        header_text = get_block_text(week_blocks[0])
        label_match = re.search(r"\[(\d{4}년\s*\d+월\s*\d+주차)\]", header_text)
        label       = label_match.group(1).replace(" ", "") if label_match else header_text.strip()

        # ── 날짜 범위 + 총건수 추출 (heading_2 직후 첫 번째 유효 paragraph)
        range_text = ""
        total      = 0
        for block in week_blocks[1:6]:
            if block.get("type") == "paragraph":
                t = get_block_text(block).strip()
                if t and "자동" not in t:
                    m_total    = re.search(r"총\s*(\d+)건", t)
                    total      = int(m_total.group(1)) if m_total else 0
                    range_text = re.sub(r"\s*총\s*\d+건.*", "", t).strip()
                    break

        # ── 섹션별 bulleted_list_item 수집
        section_map = {
            "고객사별 접촉 현황":  "clients",
            "세종본부 연관 기회":  "opps",
            "액션 필요 건":        "actions",
            "주목할 이슈·리스크":  "risks",
        }
        collected = {v: [] for v in section_map.values()}
        cur_key   = None
        generated = ""

        for block in week_blocks[1:]:
            btype = block.get("type", "")
            text  = get_block_text(block).strip()

            if btype == "heading_3":
                # 섹션 전환: 텍스트에 해당 키워드 포함 여부로 판단
                cur_key = None
                for sec_label, sec_key in section_map.items():
                    if sec_label in text:
                        cur_key = sec_key
                        break

            elif btype == "bulleted_list_item" and cur_key is not None and text:
                collected[cur_key].append(text)

            elif btype == "paragraph" and "자동 생성" in text and text:
                # 자동 생성 푸터 문자열 저장 후 섹션 수집 중단
                generated = text
                cur_key   = None

        weeks.append({
            "label":     label,
            "range":     range_text,
            "total":     total,
            "clients":   collected["clients"],
            "opps":      collected["opps"],
            "actions":   collected["actions"],
            "risks":     collected["risks"],
            "generated": generated,
        })

    return weeks


def get_sales_data(force_refresh: bool = False) -> list:
    """영업보고 전체 주차 데이터를 반환합니다. CACHE_TTL 재사용."""
    global _sales_cache
    now = time.time()

    if (not force_refresh
            and _sales_cache["data"] is not None
            and now - _sales_cache["timestamp"] < CACHE_TTL):
        print("[CACHE] 영업보고 캐시 반환")
        return _sales_cache["data"]

    print("[API] 영업보고 데이터 새로 가져오기")
    top_blocks           = fetch_block_children(SALES_PAGE_ID)
    data                 = parse_sales_blocks(top_blocks)
    _sales_cache["data"]      = data
    _sales_cache["timestamp"] = now
    return data


@app.get("/api/sales")
def api_sales(refresh: int = 0):
    """Phase 7: 영업보고 — 전체 주차 데이터 최신순 반환."""
    try:
        return get_sales_data(force_refresh=bool(refresh))
    except Exception as e:
        return {"error": str(e), "data": []}


# ========================================
# PHASE 8 — RISK 알림
# ========================================

def parse_risk_blocks(top_blocks: list) -> list:
    """heading_2 기준으로 주차 분리 후 각 주차 리스크 데이터 파싱, 최신순 반환."""
    # heading_2 중 "[20XX년" 으로 시작하는 블록만 주차 구분자로 인식
    week_positions = [
        i for i, b in enumerate(top_blocks)
        if b.get("type") == "heading_2"
        and re.search(r"\[\d{4}년", get_block_text(b))
    ]
    weeks = []
    for idx, start in enumerate(week_positions):
        end = week_positions[idx + 1] if idx + 1 < len(week_positions) else len(top_blocks)
        week_blocks = top_blocks[start:end]

        # 주차 라벨 파싱 (예: "2026년 5월 4주차")
        header_text = get_block_text(week_blocks[0])
        label_match = re.search(r"\[(\d{4}년\s*\d+월\s*\d+주차)\]", header_text)
        label = label_match.group(1).replace(" ", "") if label_match else header_text.strip()

        # 날짜 기준 / 분석 범위 파싱 (heading_2 바로 다음 paragraph)
        base_date = ""
        scope = ""
        for block in week_blocks[1:5]:
            if block.get("type") == "paragraph":
                t = get_block_text(block).strip()
                if t and "자동" not in t and "특이 리스크" not in t:
                    # "05/21 기준 | 분석 범위: 노션 3주차 + 슬랙 21일" 형태
                    parts = t.split("|")
                    base_date = parts[0].strip() if parts else ""
                    if len(parts) >= 2:
                        scope_raw = parts[1].strip()
                        # "분석 범위: " 접두어 제거
                        scope = re.sub(r"^분석\s*범위\s*:\s*", "", scope_raw).strip()
                    break

        # 섹션 키워드 → dict 키 매핑
        section_map = {
            "장기 미해결": "long_issues",
            "악화 감지":   "worsening",
            "보고 누락":   "missing",
            "일정 리스크": "schedule",
            "집중 확인":   "focus",
        }
        collected = {v: [] for v in section_map.values()}
        cur_key = None
        generated = ""
        no_risk = False

        for block in week_blocks[1:]:
            btype = block.get("type", "")
            text  = get_block_text(block).strip()

            if btype == "heading_3":
                # 섹션 식별
                cur_key = None
                for sec_label, sec_key in section_map.items():
                    if sec_label in text:
                        cur_key = sec_key
                        break
            elif btype == "bulleted_list_item" and cur_key is not None and text:
                collected[cur_key].append(text)
            elif btype == "paragraph":
                if "자동 생성" in text and text:
                    generated = text
                    cur_key = None
                elif "특이 리스크 없음" in text:
                    no_risk = True

        weeks.append({
            "label":       label,
            "base_date":   base_date,
            "scope":       scope,
            "no_risk":     no_risk,
            "long_issues": collected["long_issues"],
            "worsening":   collected["worsening"],
            "missing":     collected["missing"],
            "schedule":    collected["schedule"],
            "focus":       collected["focus"],
            "generated":   generated,
        })
    return weeks


def get_risk_data(force_refresh: bool = False) -> list:
    """RISK 알림 전체 주차 데이터를 반환합니다. CACHE_TTL 재사용."""
    global _risk_cache
    now = time.time()
    if (not force_refresh and _risk_cache["data"] is not None
            and now - _risk_cache["timestamp"] < CACHE_TTL):
        return _risk_cache["data"]
    top_blocks = fetch_block_children(RISK_PAGE_ID)
    data = parse_risk_blocks(top_blocks)
    _risk_cache["data"] = data
    _risk_cache["timestamp"] = now
    return data


@app.get("/api/risk")
def api_risk(refresh: int = 0):
    """Phase 8: RISK 알림 — 전체 주차 데이터 최신순 반환."""
    try:
        return get_risk_data(force_refresh=bool(refresh))
    except Exception as e:
        return {"error": str(e), "data": []}


# ========================================
# PHASE 9 — 세미나 알림
# ========================================

def parse_event_group(section_blocks: list) -> list:
    """
    heading_3 아래 paragraph(행사명) + bulleted_list(날짜/장소/링크) 묶음을 파싱한다.
    반환: [{"name": ..., "date": ..., "place": ..., "link": ...}, ...]
    """
    events = []
    cur = None  # 현재 수집 중인 행사 딕셔너리

    for block in section_blocks:
        btype = block.get("type", "")
        text  = get_block_text(block).strip()

        if btype == "paragraph" and text and "자동 생성" not in text:
            # 새 행사 시작 — 이전 행사 저장 후 새 딕셔너리 생성
            if cur is not None:
                events.append(cur)
            cur = {"name": text, "date": "", "place": "", "link": ""}

        elif btype == "bulleted_list_item" and cur is not None and text:
            # 키워드로 속성 분류
            if text.startswith("날짜:"):
                cur["date"] = text[3:].strip()
            elif text.startswith("장소:"):
                cur["place"] = text[3:].strip()
            elif text.startswith("링크:"):
                cur["link"] = text[3:].strip()

    # 마지막 행사 저장
    if cur is not None:
        events.append(cur)

    return events


def parse_seminar_blocks(top_blocks: list) -> list:
    """heading_2 기준으로 주차 분리 후 각 주차 세미나 데이터 파싱, 최신순 반환."""
    # heading_2 중 "[20XX년" 으로 시작하는 블록만 주차 구분자로 인식
    week_positions = [
        i for i, b in enumerate(top_blocks)
        if b.get("type") == "heading_2"
        and re.search(r"\[\d{4}년", get_block_text(b))
    ]
    weeks = []
    for idx, start in enumerate(week_positions):
        end = week_positions[idx + 1] if idx + 1 < len(week_positions) else len(top_blocks)
        week_blocks = top_blocks[start:end]

        # 주차 라벨 파싱 (예: "2026년 5월 4주차")
        header_text = get_block_text(week_blocks[0])
        label_match = re.search(r"\[(\d{4}년\s*\d+월\s*\d+주차)\]", header_text)
        label = label_match.group(1).replace(" ", "") if label_match else header_text.strip()

        # 기준일·건수 파싱 + 신규 없음 감지 (heading_2 바로 다음 paragraph)
        base_date = ""
        count = 0
        no_event = False
        generated = ""

        for block in week_blocks[1:5]:
            if block.get("type") == "paragraph":
                t = get_block_text(block).strip()
                if not t:
                    continue
                if "자동 생성" in t:
                    generated = t
                    break
                if "신규 행사 없음" in t:
                    no_event = True
                    break
                # "05/21 기준 · 신규 N건" 형태 파싱
                m_base = re.search(r"(\d{1,2}/\d{1,2}\s*기준)", t)
                m_cnt  = re.search(r"신규\s*(\d+)건", t)
                if m_base:
                    base_date = m_base.group(1)
                if m_cnt:
                    count = int(m_cnt.group(1))
                break

        # 섹션 분류: 근거리 행사 vs 기타 행사
        section_map = {
            "근거리 행사": "nearby",
            "기타 행사":   "others",
        }
        section_blocks_map = {v: [] for v in section_map.values()}
        cur_key = None

        for block in week_blocks[1:]:
            btype = block.get("type", "")
            text  = get_block_text(block).strip()

            if btype == "heading_3":
                cur_key = None
                for sec_label, sec_key in section_map.items():
                    if sec_label in text:
                        cur_key = sec_key
                        break
            elif btype == "paragraph" and "자동 생성" in text and text:
                generated = text
                cur_key = None
            elif cur_key is not None:
                # 섹션 내 블록 수집 (paragraph + bulleted_list_item)
                if btype in ("paragraph", "bulleted_list_item"):
                    section_blocks_map[cur_key].append(block)

        # 섹션별 행사 파싱
        nearby = parse_event_group(section_blocks_map["nearby"])
        others = parse_event_group(section_blocks_map["others"])

        weeks.append({
            "label":     label,
            "base_date": base_date,
            "count":     count,
            "no_event":  no_event,
            "nearby":    nearby,
            "others":    others,
            "generated": generated,
        })
    return weeks


def get_seminar_data(force_refresh: bool = False) -> list:
    """세미나 알림 전체 주차 데이터를 반환합니다. CACHE_TTL 재사용."""
    global _seminar_cache
    now = time.time()
    if (not force_refresh and _seminar_cache["data"] is not None
            and now - _seminar_cache["timestamp"] < CACHE_TTL):
        return _seminar_cache["data"]
    top_blocks = fetch_block_children(SEMINAR_PAGE_ID)
    data = parse_seminar_blocks(top_blocks)
    _seminar_cache["data"] = data
    _seminar_cache["timestamp"] = now
    return data


@app.get("/api/seminar")
def api_seminar(refresh: int = 0):
    """Phase 9: 세미나 알림 — 전체 주차 데이터 최신순 반환."""
    try:
        return get_seminar_data(force_refresh=bool(refresh))
    except Exception as e:
        return {"error": str(e), "data": []}


# ========================================
# PHASE 11 — 손익현황
# ========================================

PROFIT_LOSS_DB_ID = "f162d5cb-bf3b-4ed7-8d8a-0f2e20187d15"
_profit_loss_cache = {"records": None, "timestamp": 0}


def fetch_profit_loss_db(force_refresh: bool = False) -> list:
    """노션 DB에서 손익현황 연도별 데이터를 조회합니다."""
    global _profit_loss_cache
    now = time.time()

    if (not force_refresh
            and _profit_loss_cache["records"] is not None
            and now - _profit_loss_cache["timestamp"] < CACHE_TTL):
        return _profit_loss_cache["records"]

    url     = f"https://api.notion.com/v1/databases/{PROFIT_LOSS_DB_ID}/query"
    records = []
    body    = {"page_size": 100}

    while True:
        try:
            res = requests.post(url, headers=NOTION_HEADERS, json=body, timeout=10)
            res.raise_for_status()
            data = res.json()
            for page in data.get("results", []):
                record = parse_pl_record(page)
                if record:
                    records.append(record)
            if not data.get("has_more"):
                break
            body["start_cursor"] = data["next_cursor"]
        except Exception as e:
            print(f"[ERROR] 손익현황 DB 조회 실패: {e}")
            break

    _profit_loss_cache = {"records": records, "timestamp": now}
    return records


def parse_pl_record(page: dict) -> dict | None:
    """노션 DB 손익현황 레코드를 파싱합니다. 속성명을 한글 그대로 사용."""
    props = page.get("properties", {})

    def get_title(key):
        p = props.get(key, {})
        if p.get("type") == "title":
            return "".join(t["plain_text"] for t in p.get("title", []))
        return ""

    def get_num(key):
        p = props.get(key, {})
        if p.get("type") == "number":
            return p.get("number")
        return None

    def get_formula(key):
        p = props.get(key, {})
        if p.get("type") == "formula":
            f = p.get("formula", {})
            if f.get("type") == "number":
                return f.get("number")
        return None

    year_title = get_title("연도")
    if not year_title:
        return None

    record = {"연도": year_title}

    number_keys = [
        "전사_인력", "수행조직_인력", "세종_인력",
        "전사_매출", "수행조직_매출", "세종_매출",
        "전사_매출원가", "수행조직_매출원가", "세종_매출원가",
        "전사_매출이익", "수행조직_매출이익", "세종_매출이익",
        "전사_판관비", "수행조직_판관비", "세종_판관비",
        "전사_영업이익", "수행조직_영업이익", "세종_영업이익",
    ]
    for key in number_keys:
        record[key] = get_num(key)

    formula_keys = [
        "인력비중_수행조직대비", "인력비중_전사대비",
        "매출비중_수행조직대비", "매출비중_전사대비",
        "매출원가비중_수행조직대비", "매출원가비중_전사대비",
        "매출이익비중_수행조직대비", "매출이익비중_전사대비",
        "전사_매출이익률", "수행조직_매출이익률", "세종_매출이익률",
        "전사_영업이익률", "수행조직_영업이익률", "세종_영업이익률",
        "수행인력대비_매출_금액", "수행인력대비_매출_pp",
        "수행인력대비_매출원가_금액", "수행인력대비_매출원가_pp",
        "수행인력대비_매출이익_금액", "수행인력대비_매출이익_pp",
        "수행인력대비_매출_전사_금액", "수행인력대비_매출_전사_pp",
        "수행인력대비_매출원가_전사_금액", "수행인력대비_매출원가_전사_pp",
        "수행인력대비_매출이익_전사_금액", "수행인력대비_매출이익_전사_pp",
    ]
    for key in formula_keys:
        record[key] = get_formula(key)

    return record


def get_profit_loss_data(force_refresh: bool = False) -> dict:
    """손익현황 데이터를 연도 오름차순으로 정렬하여 반환합니다."""
    records = fetch_profit_loss_db(force_refresh)
    records.sort(key=lambda r: r.get("연도", ""))
    return {"success": True, "data": records}


@app.get("/api/profit-loss")
def api_profit_loss(refresh: int = 0):
    """Phase 11: 손익현황."""
    try:
        return get_profit_loss_data(force_refresh=bool(refresh))
    except Exception as e:
        return {"success": False, "error": str(e), "data": []}


# ========================================
# PHASE DL — 다운로드 엔드포인트 (기존 코드 수정 없음)
# ========================================

from fastapi.responses import StreamingResponse
from urllib.parse import quote as _url_quote

def _docx_resp(content: bytes, filename: str) -> StreamingResponse:
    """워드 파일을 StreamingResponse로 반환합니다."""
    encoded = _url_quote(filename)
    return StreamingResponse(
        iter([content]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )


def _pdf_resp(content: bytes, filename: str) -> StreamingResponse:
    """PDF 파일을 StreamingResponse로 반환합니다."""
    encoded = _url_quote(filename)
    return StreamingResponse(
        iter([content]),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )


def _fmt_eok(val) -> str:
    """원 단위 숫자 → '123.4억' 형식."""
    if val is None:
        return "—"
    try:
        return f"{float(val) / 1e8:.1f}억"
    except Exception:
        return "—"


def _fmt_num(val) -> str:
    """정수 → 콤마 포맷."""
    if val is None:
        return "—"
    try:
        return f"{int(val):,}"
    except Exception:
        return str(val)


# ── 대시보드 다운로드 ──────────────────────────────────────────────────────

@app.get("/api/dashboard/download")
def api_dashboard_download():
    """대시보드 현황 워드 다운로드."""
    try:
        from report_template import (
            new_document, make_title_page, make_heading1,
            make_kpi_table, make_data_table, embed_chart_image,
            add_footnote, setup_header_footer, doc_to_bytes,
            make_chart_bar, add_bullet, NAVY_HEX, GOLD_HEX, RED_HEX,
        )
        from datetime import datetime

        projects_raw = get_notion_projects()
        risk_raw     = get_risk_data()
        sales_raw    = get_sales_data()
        org_raw      = get_org_data()

        active   = [p for p in projects_raw if p.get("status") != "준비중"]
        ready    = [p for p in projects_raw if p.get("status") == "준비중"]
        tot_rev  = sum(p.get("revenue") or 0 for p in active)
        staff    = org_raw.get("staff", [])
        headcount = len([s for s in staff if s.get("type") != "협력"])

        latest_sales   = sales_raw[0]   if sales_raw   else {}
        latest_risk    = risk_raw[0]    if risk_raw    else {}

        doc = new_document()
        setup_header_footer(doc, "대시보드")
        make_title_page(doc, "대시보드 현황 보고",
                        subtitle=f"생성일: {datetime.now().strftime('%Y.%m.%d')}")

        # KPI 5개
        make_heading1(doc, "핵심 지표")
        make_kpi_table(doc, [
            ("수행 사업",  f"{len(active)}건",                          "",                   "navy"),
            ("총 매출액",  _fmt_eok(tot_rev),                           "준비중 제외",        "gold"),
            ("팀 인원",    f"{headcount}명",                            "협력 제외",          "navy"),
            ("이번주 영업", f"{latest_sales.get('total', 0)}건",         latest_sales.get('label', ''),   "blue"),
            ("준비중 사업", f"{len(ready)}건",                           "",                  "gray"),
        ])

        # 수행 사업 현황 표
        make_heading1(doc, "수행 사업 현황")
        rows = [
            [p.get("name", ""), p.get("client", ""), p.get("pm", ""),
             _fmt_eok(p.get("revenue")), p.get("status", ""), p.get("period", "")]
            for p in active
        ]
        make_data_table(doc, ["사업명", "고객사", "PM", "매출액", "상태", "기간"], rows)

        # 사업별 매출 차트
        if active:
            chart_data = sorted(active, key=lambda p: p.get("revenue") or 0, reverse=True)[:10]
            labels = [p["name"][:8] for p in chart_data]
            values = [round((p.get("revenue") or 0) / 1e8, 1) for p in chart_data]
            img = make_chart_bar(labels, values, title="사업별 매출 구성 (억원)", bar_color=NAVY_HEX)
            embed_chart_image(doc, img, width_cm=14)
            add_footnote(doc, "출처: 노션 사업현황 DB  ※ 준비중 제외")

        # RISK 요약 (최대 5건)
        make_heading1(doc, "RISK 알림 요약")
        if latest_risk and not latest_risk.get("no_risk"):
            label = latest_risk.get("label", "")
            add_footnote(doc, f"기준: {label}")
            risk_items = (
                latest_risk.get("long_issues", []) +
                latest_risk.get("worsening",   []) +
                latest_risk.get("schedule",    [])
            )
            for item in risk_items[:5]:
                add_bullet(doc, item)
        else:
            add_bullet(doc, "특이 리스크 없음")

        # 영업보고 요약
        make_heading1(doc, "영업보고 요약")
        if latest_sales:
            add_footnote(doc, f"기준: {latest_sales.get('label', '')}  {latest_sales.get('range', '')}")
            for item in latest_sales.get("clients", [])[:5]:
                add_bullet(doc, item)

        filename = f"대시보드_{datetime.now().strftime('%Y%m%d')}.docx"
        return _docx_resp(doc_to_bytes(doc), filename)

    except Exception as e:
        return {"error": str(e)}


# ── 사업현황 다운로드 ──────────────────────────────────────────────────────

@app.get("/api/projects/download")
def api_projects_download():
    """사업현황 워드 다운로드."""
    try:
        from report_template import (
            new_document, make_title_page, make_heading1,
            make_kpi_table, make_data_table, embed_chart_image,
            add_footnote, setup_header_footer, doc_to_bytes,
            make_chart_bar, NAVY_HEX, GOLD_HEX,
        )
        from datetime import datetime

        projects = get_notion_projects()
        active   = [p for p in projects if p.get("status") != "준비중"]
        ready    = [p for p in projects if p.get("status") == "준비중"]
        tot_rev  = sum(p.get("revenue") or 0 for p in active)

        doc = new_document()
        setup_header_footer(doc, "사업현황")
        make_title_page(doc, "사업현황 보고",
                        subtitle=f"생성일: {datetime.now().strftime('%Y.%m.%d')}")

        # KPI
        make_heading1(doc, "핵심 지표")
        make_kpi_table(doc, [
            ("수행 사업",  f"{len(active)}건",     "",           "navy"),
            ("총 매출액",  _fmt_eok(tot_rev),      "준비중 제외","gold"),
            ("준비중 사업", f"{len(ready)}건",     "",           "blue"),
        ])

        # 매출 구성 차트
        if active:
            sorted_p = sorted(active, key=lambda p: p.get("revenue") or 0, reverse=True)
            labels = [p["name"][:10] for p in sorted_p]
            values = [round((p.get("revenue") or 0) / 1e8, 1) for p in sorted_p]
            img = make_chart_bar(labels, values, title="사업별 매출 구성 (억원)", bar_color=NAVY_HEX)
            embed_chart_image(doc, img, width_cm=14)
            add_footnote(doc, "출처: 노션 사업현황 DB  ※ 준비중 제외")

        # 수행 사업 표
        make_heading1(doc, "수행 사업 목록")
        rows = [
            [p.get("name",""), p.get("client",""), p.get("pm",""),
             _fmt_eok(p.get("revenue")), p.get("period",""), p.get("status","")]
            for p in active
        ]
        make_data_table(doc, ["사업명","고객사","PM","매출액","기간","상태"], rows)

        # 준비중 사업
        if ready:
            make_heading1(doc, "준비중 사업")
            rows2 = [
                [p.get("name",""), p.get("client",""), p.get("pm",""),
                 _fmt_eok(p.get("revenue")), p.get("period","")]
                for p in ready
            ]
            make_data_table(doc, ["사업명","고객사","PM","예상매출","기간"], rows2)

        filename = f"사업현황_{datetime.now().strftime('%Y%m%d')}.docx"
        return _docx_resp(doc_to_bytes(doc), filename)

    except Exception as e:
        return {"error": str(e)}


# ── 주간회의 다운로드 ──────────────────────────────────────────────────────

@app.get("/api/weekly/download")
def api_weekly_download(week: int = 0):
    """PM 주간회의 워드 다운로드. ?week=N 으로 주차 지정."""
    try:
        from report_template import (
            new_document, make_title_page, make_heading1,
            make_data_table, add_footnote, add_bullet,
            setup_header_footer, doc_to_bytes,
        )
        from docx.shared import Pt
        from datetime import datetime

        data = get_weekly_data(week_index=week)
        if not data.get("success"):
            return {"error": "주간회의 데이터 없음"}

        d_inner    = data.get("data", {})
        week_label = d_inner.get("week_title", f"{week+1}주차")
        projects   = d_inner.get("projects", [])
        analysis   = d_inner.get("analysis", [])

        doc = new_document()
        setup_header_footer(doc, "PM 주간회의")
        make_title_page(doc, f"{week_label} 주간회의록",
                        subtitle=f"생성일: {datetime.now().strftime('%Y.%m.%d')}")

        # 프로젝트별 진행사항
        make_heading1(doc, "프로젝트별 현황")
        for proj in projects:
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(8)
            r = p_title.add_run(proj.get("name", ""))
            r.bold      = True
            r.font.size = Pt(11)

            # 실제 필드명: progress / special / issues
            for section_key, section_label in [
                ("progress", "진행사항"),
                ("special",  "특이사항"),
                ("issues",   "이슈사항"),
            ]:
                items = proj.get(section_key, [])
                if items:
                    lbl_p = doc.add_paragraph()
                    lbl_r = lbl_p.add_run(f"[{section_label}]")
                    lbl_r.bold      = True
                    lbl_r.font.size = Pt(10)
                    for item in items:
                        add_bullet(doc, item)

        # 분석 섹션
        if analysis:
            make_heading1(doc, "분석")
            for item in analysis:
                add_bullet(doc, item)

        add_footnote(doc, f"출처: 노션 PM 주간회의  ※ {week_label} 기준")
        filename = f"주간회의_{week_label.replace(' ', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return _docx_resp(doc_to_bytes(doc), filename)

    except Exception as e:
        return {"error": str(e)}


# ── 월간보고 다운로드 ──────────────────────────────────────────────────────

@app.get("/api/monthly/download")
def api_monthly_download(month: int = 0):
    """회사 월간보고 워드 다운로드. ?month=N 으로 월 지정."""
    try:
        from report_template import (
            new_document, make_title_page, make_heading1,
            make_kpi_table, make_data_table, add_footnote, add_bullet,
            setup_header_footer, doc_to_bytes,
        )
        from datetime import datetime

        data = get_monthly_data_company(month_index=month)
        if not data.get("success"):
            return {"error": "월간보고 데이터 없음"}

        d           = data.get("data", {})
        month_title = d.get("month_title", "")
        kpi         = d.get("kpi", {})
        workforce   = d.get("workforce", [])
        issue_table = d.get("issue_table", [])
        meetings    = d.get("meetings", [])
        analysis    = d.get("analysis", [])

        doc = new_document()
        setup_header_footer(doc, "월간보고")
        make_title_page(doc, f"{month_title} 월간보고",
                        subtitle=f"생성일: {datetime.now().strftime('%Y.%m.%d')}")

        # KPI
        make_heading1(doc, "핵심 지표")
        make_kpi_table(doc, [
            ("진행 프로젝트",  f"{kpi.get('total', 0)}건",   "", "navy"),
            ("이슈 프로젝트",  f"{kpi.get('issue', 0)}건",   "", "red"),
            ("신규 미팅",      f"{kpi.get('meeting', 0)}건", "", "blue"),
        ])

        # 인력투입현황
        if workforce:
            make_heading1(doc, "인력투입현황")
            rows = [[w.get("label",""), w.get("value","")] for w in workforce]
            make_data_table(doc, ["구분","투입인원"], rows)

        # 이슈사항
        if issue_table:
            make_heading1(doc, "이슈사항")
            rows = [
                [i.get("client",""), i.get("project",""), i.get("budget",""),
                 i.get("period",""), i.get("status",""), i.get("issue","")]
                for i in issue_table
            ]
            make_data_table(doc, ["고객사","사업명","예산","기간","단계","이슈"], rows)

        # 분석
        if analysis:
            make_heading1(doc, "분석")
            for item in analysis:
                add_bullet(doc, item)

        add_footnote(doc, f"출처: 노션 회사 월간보고  ※ {month_title} 기준")
        filename = f"월간보고_{month_title.replace(' ','')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return _docx_resp(doc_to_bytes(doc), filename)

    except Exception as e:
        return {"error": str(e)}


# ── 사업히스토리 다운로드 ─────────────────────────────────────────────────

@app.get("/api/history/download")
def api_history_download(year: str = ""):
    """사업히스토리 워드 다운로드. ?year=YYYY 로 연도 지정."""
    try:
        from report_template import (
            new_document, make_title_page, make_heading1,
            make_kpi_table, make_data_table, embed_chart_image,
            add_footnote, setup_header_footer, doc_to_bytes,
            make_chart_bar_grouped, NAVY_HEX, GOLD_HEX, BLUE_HEX,
        )
        from datetime import datetime

        data    = get_history_data()
        summary = data.get("data", {}).get("summary", [])
        years   = data.get("data", {}).get("years", [])

        # 운영매출 추이 차트 (전체 연도)
        doc = new_document()
        setup_header_footer(doc, "사업히스토리")
        make_title_page(doc, "세종사업히스토리",
                        subtitle=f"생성일: {datetime.now().strftime('%Y.%m.%d')}")

        if summary:
            x_labels = [f"{s['year']}년" for s in summary]
            ops_vals = []
            non_vals = []
            for s in summary:
                yr_num   = s["year"]
                yr_entry = next((y for y in years if y["year"] == yr_num), None)
                if yr_entry:
                    chart = yr_entry.get("chart", {})
                    ops_vals.append(round(chart.get("ops", 0) / 1e8, 1))
                    non_vals.append(round(chart.get("non_ops", 0) / 1e8, 1))
                else:
                    ops_vals.append(round(s.get("revenue", 0) / 1e8, 1))
                    non_vals.append(0)

            img = make_chart_bar_grouped(
                x_labels,
                [("운영매출", ops_vals, NAVY_HEX), ("운영외매출", non_vals, GOLD_HEX)],
                title="연도별 매출 추이 (억원)",
                ylabel="억원",
            )
            embed_chart_image(doc, img, width_cm=14)
            add_footnote(doc, "출처: 노션 사업히스토리 DB  ※ 내부원가현행화 기준")

        # 연도별 요약 표
        make_heading1(doc, "연도별 사업 현황 요약")
        rows = [
            [
                f"{s['year']}년",
                _fmt_eok(s.get("revenue")),
                _fmt_eok(s.get("gross_profit")),
                s.get("gross_rate", "—"),
                _fmt_eok(s.get("op_profit")),
                s.get("op_rate", "—"),
            ]
            for s in summary
        ]
        make_data_table(doc,
                        ["연도","매출액","매출이익","매출이익률","영업이익","영업이익률"],
                        rows)

        # 선택 연도 프로젝트 상세
        target_year = year if year else (summary[-1]["year"] if summary else "")
        yr_entry = next((y for y in years if y["year"] == target_year), None)
        if yr_entry and yr_entry.get("projects"):
            make_heading1(doc, f"{target_year}년 프로젝트별 상세")
            p_rows = [
                [
                    p.get("name",""), p.get("project_type",""), p.get("client",""),
                    p.get("period",""), _fmt_eok(p.get("revenue")),
                    _fmt_eok(p.get("gross_profit")), p.get("gross_rate","—"),
                ]
                for p in yr_entry["projects"]
                if not p.get("is_excluded")
            ]
            make_data_table(doc,
                            ["사업명","구분","고객사","기간","매출액","매출이익","이익률"],
                            p_rows)

        filename = f"사업히스토리_{datetime.now().strftime('%Y%m%d')}.docx"
        return _docx_resp(doc_to_bytes(doc), filename)

    except Exception as e:
        return {"error": str(e)}


# ── 영업보고 다운로드 ──────────────────────────────────────────────────────

@app.get("/api/sales/download")
def api_sales_download(week: int = 0):
    """영업보고 워드 다운로드. ?week=N 으로 주차 지정 (0=최신)."""
    try:
        from report_template import (
            new_document, make_title_page, make_heading1,
            make_data_table, add_footnote, add_bullet,
            setup_header_footer, doc_to_bytes,
        )
        from datetime import datetime

        all_weeks = get_sales_data()
        if not all_weeks or week >= len(all_weeks):
            return {"error": "영업보고 데이터 없음"}

        w = all_weeks[week]

        doc = new_document()
        setup_header_footer(doc, "영업보고")
        make_title_page(doc, f"{w.get('label','')} 영업보고 분석",
                        subtitle=f"{w.get('range','')}  ※ 총 {w.get('total',0)}건\n생성일: {datetime.now().strftime('%Y.%m.%d')}")

        # 고객사별 접촉 현황
        if w.get("clients"):
            make_heading1(doc, "고객사별 접촉 현황")
            for item in w["clients"]:
                add_bullet(doc, item)

        # 세종본부 연관 기회
        if w.get("opps"):
            make_heading1(doc, "세종본부 연관 기회")
            for item in w["opps"]:
                add_bullet(doc, item)

        # 액션 필요 건
        if w.get("actions"):
            make_heading1(doc, "액션 필요 건")
            for item in w["actions"]:
                add_bullet(doc, item)

        # 이슈·리스크
        if w.get("risks"):
            make_heading1(doc, "주목할 이슈·리스크")
            for item in w["risks"]:
                add_bullet(doc, item)

        add_footnote(doc, f"출처: 노션 영업보고  ※ {w.get('label','')} 기준")
        filename = f"영업보고_{w.get('label','').replace(' ','')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return _docx_resp(doc_to_bytes(doc), filename)

    except Exception as e:
        return {"error": str(e)}


# ── 손익현황 다운로드 ──────────────────────────────────────────────────────

@app.get("/api/profit-loss/download")
def api_profit_loss_download():
    """손익현황 워드 다운로드."""
    try:
        from report_template import (
            new_document, make_title_page, make_heading1,
            make_kpi_table, make_data_table, embed_chart_image,
            add_footnote, setup_header_footer, doc_to_bytes,
            make_chart_bar_grouped, NAVY_HEX, GOLD_HEX, BLUE_HEX, RED_HEX,
        )
        from datetime import datetime

        result  = get_profit_loss_data()
        records = result.get("data", [])

        doc = new_document()
        setup_header_footer(doc, "손익현황")
        make_title_page(doc, "손익현황 보고",
                        subtitle=f"생성일: {datetime.now().strftime('%Y.%m.%d')}")

        # 최신 연도 KPI
        if records:
            latest = records[-1]
            make_heading1(doc, f"{latest.get('연도','')} 핵심 지표")
            gp_rate = latest.get("세종_매출이익률")
            op_rate = latest.get("세종_영업이익률")
            make_kpi_table(doc, [
                ("세종 매출",    _fmt_eok(latest.get("세종_매출")),    "", "navy"),
                ("세종 매출이익", _fmt_eok(latest.get("세종_매출이익")), "", "gold"),
                ("매출이익률",   f"{gp_rate*100:.1f}%" if gp_rate else "—", "", "blue"),
                ("영업이익",     _fmt_eok(latest.get("세종_영업이익")), "", "navy"),
                ("영업이익률",   f"{op_rate*100:.1f}%" if op_rate else "—", "", "blue"),
            ])

        # 매출/이익 추이 차트
        if records:
            x_labels = [r.get("연도","").replace("년","") for r in records]
            rev_vals = [round((r.get("세종_매출") or 0) / 1e8, 1) for r in records]
            gp_vals  = [round((r.get("세종_매출이익") or 0) / 1e8, 1) for r in records]
            img = make_chart_bar_grouped(
                x_labels,
                [("매출", rev_vals, NAVY_HEX), ("매출이익", gp_vals, GOLD_HEX)],
                title="세종개발본부 매출·이익 추이 (억원)",
                ylabel="억원",
            )
            embed_chart_image(doc, img, width_cm=14)
            add_footnote(doc, "출처: 노션 손익현황 DB  ※ 세종개발본부 기준")

        # 연도별 요약 표
        make_heading1(doc, "연도별 손익 요약")
        rows = []
        for r in records:
            gp_r  = r.get("세종_매출이익률")
            op_r  = r.get("세종_영업이익률")
            rows.append([
                r.get("연도",""),
                _fmt_eok(r.get("세종_매출")),
                _fmt_eok(r.get("세종_매출원가")),
                _fmt_eok(r.get("세종_매출이익")),
                f"{gp_r*100:.1f}%" if gp_r else "—",
                _fmt_eok(r.get("세종_영업이익")),
                f"{op_r*100:.1f}%" if op_r else "—",
                f"{r.get('세종_인력') or 0}명",
            ])
        make_data_table(doc,
                        ["연도","매출","원가","매출이익","이익률","영업이익","영업이익률","인력"],
                        rows)

        filename = f"손익현황_{datetime.now().strftime('%Y%m%d')}.docx"
        return _docx_resp(doc_to_bytes(doc), filename)

    except Exception as e:
        return {"error": str(e)}


# ── 조직도 PDF 다운로드 ────────────────────────────────────────────────────

@app.get("/api/org/download-pdf")
def api_org_download_pdf():
    """조직도 인물 명단 PDF 다운로드."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER
        import io as _io
        from datetime import datetime

        # 한글 폰트 등록
        font_name = "Malgun"
        font_paths = [
            r"C:\Windows\Fonts\malgun.ttf",
            r"C:\Windows\Fonts\NanumGothic.ttf",
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        ]
        registered = False
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont(font_name, fp))
                    registered = True
                    break
                except Exception:
                    continue

        if not registered:
            font_name = "Helvetica"  # 폰트 없으면 영문 폰트로 폴백

        org = get_org_data()
        staff    = org.get("staff", [])
        projects = org.get("projects", [])

        buf = _io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2*cm,  bottomMargin=2*cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title", fontName=font_name, fontSize=18,
            textColor=colors.HexColor("#1B2A4A"),
            alignment=TA_CENTER, spaceAfter=12,
        )
        h1_style = ParagraphStyle(
            "H1", fontName=font_name, fontSize=12,
            textColor=colors.HexColor("#1B2A4A"),
            spaceBefore=12, spaceAfter=6,
        )
        cell_style = ParagraphStyle(
            "Cell", fontName=font_name, fontSize=9,
        )

        NAVY_RL = colors.HexColor("#1B2A4A")
        LGRAY_RL = colors.HexColor("#F3F4F6")

        story = []
        story.append(Paragraph("세종개발본부 조직도", title_style))
        story.append(Paragraph(
            f"생성일: {datetime.now().strftime('%Y.%m.%d')}",
            ParagraphStyle("Sub", fontName=font_name, fontSize=9,
                           textColor=colors.HexColor("#6B7280"), alignment=TA_CENTER),
        ))
        story.append(Spacer(1, 0.5*cm))

        # 사업별로 인력 분류
        proj_names = [p["name"] for p in projects] + [""]

        def _cell(text):
            return Paragraph(str(text) if text else "—", cell_style)

        for proj_name in proj_names:
            members = [s for s in staff if s.get("project") == proj_name]
            if not members:
                continue

            section_label = proj_name if proj_name else "공통/미배정"
            story.append(Paragraph(section_label, h1_style))

            header = ["이름","직급","직책","구분"]
            data   = [header] + [
                [_cell(s.get("name","")), _cell(s.get("grade","")),
                 _cell(s.get("role","")),  _cell(s.get("type",""))]
                for s in members
            ]

            col_widths = [4*cm, 3*cm, 3.5*cm, 3*cm]
            t = Table(data, colWidths=col_widths)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), NAVY_RL),
                ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                ("FONTNAME",   (0,0), (-1,-1), font_name),
                ("FONTSIZE",   (0,0), (-1,-1), 9),
                ("ALIGN",      (0,0), (-1,-1), "CENTER"),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, LGRAY_RL]),
                ("GRID",       (0,0), (-1,-1), 0.3, colors.HexColor("#E5E7EB")),
                ("TOPPADDING", (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.3*cm))

        def _footer(canvas, doc_inner):
            canvas.saveState()
            canvas.setFont(font_name, 8)
            canvas.setFillColor(colors.HexColor("#6B7280"))
            canvas.drawCentredString(A4[0]/2, 1.2*cm, "유라클 세종개발본부")
            canvas.restoreState()

        doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
        buf.seek(0)

        filename = f"조직도_{datetime.now().strftime('%Y%m%d')}.pdf"
        return _pdf_resp(buf.read(), filename)

    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8001))
    uvicorn.run(app, host="0.0.0.0", port=port)