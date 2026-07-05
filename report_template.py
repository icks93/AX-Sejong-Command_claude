"""
report_template.py — AX Sejong Command 공통 보고서 템플릿
8개 메뉴의 워드/PDF 문서가 이 모듈 함수를 공유합니다.
"""

import io
import os

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

from datetime import datetime
from urllib.parse import quote

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 색상 상수 (RGBColor: python-docx용)
NAVY       = RGBColor(0x1B, 0x2A, 0x4A)
GOLD       = RGBColor(0xB8, 0x86, 0x0B)
RED        = RGBColor(0xC5, 0x30, 0x30)
BLUE       = RGBColor(0x2B, 0x6C, 0xB0)
TEXT_COLOR = RGBColor(0x1A, 0x1F, 0x2B)
GRAY       = RGBColor(0x6B, 0x72, 0x80)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── hex 상수 (matplotlib용, XML fill용)
NAVY_HEX  = "#1B2A4A"
GOLD_HEX  = "#B8860B"
RED_HEX   = "#C53030"
BLUE_HEX  = "#2B6CB0"
GRAY_HEX  = "#6B7280"
LGRAY_HEX = "#F3F4F6"  # 연회색 배경

_font_ready = False


# ── 내부 헬퍼 ──────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str):
    """셀 배경색 설정."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _para_border_bottom(para, color_hex: str = "1B2A4A", sz: int = 6):
    """단락 하단에 구분선(border-bottom) 추가."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)


def _cell_no_margin(cell):
    """셀 내부 여백 최소화."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    "60")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


# ── 공개 API ───────────────────────────────────────────────────────────────

def setup_matplotlib_font():
    """
    matplotlib 한글 폰트 설정 (최초 1회).
    우선순위: Noto Sans CJK KR → Malgun Gothic (Windows) → NanumGothic
    """
    global _font_ready
    if _font_ready:
        return

    available = {f.name for f in fm.fontManager.ttflist}
    candidates = [
        "Noto Sans CJK KR",
        "NotoSansCJKkr",
        "Noto Sans KR",
        "NanumGothic",
        "Malgun Gothic",
    ]
    chosen = next((c for c in candidates if c in available), None)

    if not chosen:
        # 시스템 폰트 파일 직접 검색
        font_paths = [
            r"C:\Windows\Fonts\malgun.ttf",
            r"C:\Windows\Fonts\NanumGothic.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
        ]
        for fp in font_paths:
            if os.path.exists(fp):
                fm.fontManager.addfont(fp)
                prop  = fm.FontProperties(fname=fp)
                chosen = prop.get_name()
                break

    if chosen:
        matplotlib.rcParams["font.family"]       = chosen
        matplotlib.rcParams["axes.unicode_minus"] = False

    _font_ready = True


def new_document() -> Document:
    """기본 여백이 설정된 빈 Document 반환."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    return doc


def setup_header_footer(doc: Document, menu_name: str):
    """
    헤더: 우측 "{메뉴명}  |  YYYY.MM" (회색 소문자)
    푸터: 중앙 "유라클 세종개발본부" (회색)
    """
    now_str = datetime.now().strftime("%Y.%m")
    for section in doc.sections:
        # 헤더
        hp = section.header.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(f"{menu_name}  |  {now_str}")
        r.font.size      = Pt(8)
        r.font.color.rgb = GRAY

        # 푸터
        fp = section.footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = fp.add_run("유라클 세종개발본부")
        r2.font.size      = Pt(8)
        r2.font.color.rgb = GRAY


def make_title_page(doc: Document, menu_name: str, subtitle: str = ""):
    """표지: 메뉴명(28pt 네이비 굵게) + 부제(11pt 회색)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(menu_name)
    r.bold            = True
    r.font.size       = Pt(28)
    r.font.color.rgb  = NAVY

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(8)
        r2 = p2.add_run(subtitle)
        r2.font.size      = Pt(11)
        r2.font.color.rgb = GRAY

    doc.add_paragraph()


def make_heading1(doc: Document, text: str):
    """섹션 제목: 14pt 네이비 굵게 + 하단 네이비 구분선."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(14)
    r.font.color.rgb = NAVY
    _para_border_bottom(p)


def make_kpi_table(doc: Document, items: list):
    """
    KPI 타일 (3행 × n열).
    items = [(label, value, sub, color_key), ...]
    color_key: "navy" | "gold" | "red" | "blue" | None
    """
    n = len(items)
    if not n:
        return

    color_map = {"navy": NAVY, "gold": GOLD, "red": RED, "blue": BLUE}
    table = doc.add_table(rows=3, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    for ci, (label, value, sub, color_key) in enumerate(items):
        val_color = color_map.get(color_key, NAVY)
        for ri in range(3):
            _cell_bg(table.cell(ri, ci), LGRAY_HEX.lstrip("#"))
            _cell_no_margin(table.cell(ri, ci))

        # 라벨 행
        p0 = table.cell(0, ci).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(label))
        r0.font.size      = Pt(9)
        r0.font.color.rgb = GRAY

        # 값 행
        p1 = table.cell(1, ci).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(str(value))
        r1.bold           = True
        r1.font.size      = Pt(18)
        r1.font.color.rgb = val_color

        # 보조 행
        p2 = table.cell(2, ci).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub:
            r2 = p2.add_run(str(sub))
            r2.font.size      = Pt(8)
            r2.font.color.rgb = GRAY

    doc.add_paragraph()


def make_data_table(doc: Document, headers: list, rows: list):
    """
    데이터 표.
    헤더: 네이비 배경 + 흰 글씨 + 굵게
    데이터: 짝수 행 연회색 (zebra)
    첫 컬럼 좌측 정렬, 나머지 가운데 정렬
    """
    if not headers:
        return

    n_cols = len(headers)
    table  = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    # 헤더 행
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        _cell_bg(cell, "1B2A4A")
        _cell_no_margin(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        r.bold           = True
        r.font.size      = Pt(9)
        r.font.color.rgb = WHITE

    # 데이터 행
    for ri, row in enumerate(rows):
        is_even = (ri % 2 == 1)
        for ci in range(n_cols):
            cell = table.cell(ri + 1, ci)
            _cell_no_margin(cell)
            if is_even:
                _cell_bg(cell, LGRAY_HEX.lstrip("#"))
            val = row[ci] if ci < len(row) else ""
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("" if val is None else str(val))
            r.font.size = Pt(9)

    doc.add_paragraph()


def embed_chart_image(doc: Document, image_bytes: bytes, width_cm: float = 14.0):
    """PNG bytes를 본문 중앙에 삽입합니다."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(image_bytes), width=Cm(width_cm))


def add_footnote(doc: Document, text: str):
    """출처/각주: 8pt 회색 이탤릭."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.italic         = True
    r.font.size      = Pt(8)
    r.font.color.rgb = GRAY


def add_bullet(doc: Document, text: str, indent_cm: float = 0.5):
    """불릿 항목 한 줄 추가."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent_cm)
    r = p.add_run(text)
    r.font.size = Pt(10)


def doc_to_bytes(doc: Document) -> bytes:
    """Document를 bytes로 변환합니다."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── 차트 생성 헬퍼 ──────────────────────────────────────────────────────────

def make_chart_bar(
    x_labels: list,
    values: list,
    title: str = "",
    ylabel: str = "",
    bar_color: str = None,
) -> bytes:
    """단순 세로 막대 차트 → PNG bytes."""
    setup_matplotlib_font()
    bar_color = bar_color or NAVY_HEX

    fig, ax = plt.subplots(figsize=(10, 4))
    bars = ax.bar(x_labels, values, color=bar_color, width=0.5)

    # 값 라벨
    for bar, v in zip(bars, values):
        if v:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height() + max(values) * 0.01,
                str(v),
                ha="center", va="bottom", fontsize=8, color=NAVY_HEX
            )

    if title:
        ax.set_title(title, fontsize=12, color=NAVY_HEX, pad=10)
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=9, color=GRAY_HEX)

    _style_axes(ax)
    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def make_chart_bar_grouped(
    x_labels: list,
    series: list,
    title: str = "",
    ylabel: str = "",
) -> bytes:
    """
    그룹 막대 차트 → PNG bytes.
    series = [(label, values, hex_color), ...]
    """
    setup_matplotlib_font()
    n        = len(x_labels)
    n_series = len(series)
    width    = 0.6 / max(n_series, 1)

    fig, ax = plt.subplots(figsize=(10, 4))
    for i, (label, values, color) in enumerate(series):
        offset = (i - n_series / 2 + 0.5) * width
        xs     = [j + offset for j in range(n)]
        ax.bar(xs, values, width=width * 0.88, label=label, color=color)

    ax.set_xticks(range(n))
    ax.set_xticklabels(x_labels, fontsize=9, rotation=15)
    ax.tick_params(axis="y", labelsize=9)
    if title:
        ax.set_title(title, fontsize=12, color=NAVY_HEX, pad=10)
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=9, color=GRAY_HEX)
    ax.legend(fontsize=8, loc="upper left")

    _style_axes(ax)
    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _style_axes(ax):
    """공통 축 스타일."""
    ax.set_facecolor("#FAFAFA")
    ax.get_figure().patch.set_facecolor("white")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#E5E7EB")
    ax.spines["bottom"].set_color("#E5E7EB")
    ax.yaxis.grid(True, color="#E5E7EB", linewidth=0.5)
    ax.set_axisbelow(True)
